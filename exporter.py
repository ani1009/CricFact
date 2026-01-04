from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
import datetime

def export_pdf(errors):
    file = "cricket_fact_report.pdf"
    c = canvas.Canvas(file, pagesize=A4)
    t = c.beginText(40, 800)

    t.textLine("Cricket Fact Check Report")
    t.textLine(f"Date: {datetime.date.today()}")
    t.textLine("")

    for e in errors:
        if e.get("table"):
            t.textLine(
                f"[Table {e['table']} | Row {e['row']} | Column {e['column']}] {e['text']}"
            )
        else:
            t.textLine(e["text"])

    c.drawText(t)
    c.save()
    return file


def export_word(errors):
    file = "cricket_fact_report.docx"
    doc = Document()
    doc.add_heading("Cricket Fact Check Report", 1)
    doc.add_paragraph(f"Date: {datetime.date.today()}")

    for e in errors:
        if e.get("table"):
            doc.add_paragraph(
                f"[Table {e['table']} | Row {e['row']} | Column {e['column']}]\n{e['text']}"
            )
        else:
            doc.add_paragraph(e["text"])

    doc.save(file)
    return file
