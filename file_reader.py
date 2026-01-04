from docx import Document
from PyPDF2 import PdfReader

def read_word(file):
    doc = Document(file)
    content = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append({
                "text": para.text.strip(),
                "source": "paragraph"
            })

    # Tables
    for t_idx, table in enumerate(doc.tables, start=1):
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        for r_idx, row in enumerate(table.rows[1:], start=2):
            values = [cell.text.strip() for cell in row.cells]

            if len(headers) != len(values):
                continue

            name = values[0]
            for h, v in zip(headers[1:], values[1:]):
                if not v:
                    continue

                sentence = f"{name} {h} {v}"
                content.append({
                    "text": sentence,
                    "source": "table",
                    "table": t_idx,
                    "row": r_idx,
                    "column": h
                })

    return content


def read_pdf(file):
    reader = PdfReader(file)
    data = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            data.append({
                "text": text,
                "source": "pdf"
            })

    return data
